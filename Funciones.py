import moviepy.editor
import moviepy as mp
import docx
from pytube import YouTube
from moviepy.editor import *
from docx import Document
import speech_recognition as sr


#1)
def descargar_audio():
    #Sirve solamente para descargar el audio de un video de youtube
    url = input("Ingrese la URL del audio: ")
    yt = YouTube(url)
    print("La canción que usted eligió es: " + yt.title)
    print(yt.streams.filter(only_audio=True))
    tag = int(input("Ingrese el tag que quiere usar: "))
    stream = yt.streams.get_by_itag(tag)
    stream.download()
    print("¡Su audio ha sido descargado de forma exitosa!")


#2)
def descargar_video():
    #Sirve para descargar el audio y video de youtube, juntos
    url = input("Ingrese la URL del video: ")
    yt = YouTube(url)
    print("La canción que usted eligió es: " + yt.title)
    print(yt.streams.filter(progressive=True))
    tag = int(input("Ingrese el tag que quiere usar: "))
    stream = yt.streams.get_by_itag(tag)
    stream.download()
    print("¡Su video ha sido descargado de forma exitosa!")


#3)
def extraccion_audio():
    #Sirve para extraer audio de un video que se tenga
    nombre_clip = input("Ingrese el nombre del archivo completo(video, junto con la extensión de este): ")
    my_clip = mp.editor.VideoFileClip(r"{nombre}".format(nombre= nombre_clip))
    my_clip.audio.write_audiofile("Audio.mp3")
    print("El audio de su video ha sido extraído de forma exitosa")


#4
def conversor_video_a_audio():
    #Convierte un video a formato de audio
    clip = input("Ingrese el nombre del archivo completo(video, junto con la extensión de este): ")
    video = moviepy.editor.VideoFileClip(clip)
    audio = video.audio
    nombre_arhivo = input("Ingrese como quiere escribir el audio y la extensión(Ej: Audio.wav): ")
    audio.write_audiofile(nombre_arhivo)
    print("El proceso se ha llevado a cabo con éxito")


#5
def audio_a_texto():
    #Convierte los videos a texto
    r = sr.Recognizer()
    print("Recuerda que tienes que tener los documentos con el mismo nombre principal y distinto número")
    nombre_documento = input("Ingresa el nombre del documento y extension. Ej: Clip1.wav: ")
    inicio = int(input("Ingresa el número del primer documento: "))
    fin = int(input("Ingresa el número del último documento: "))

    for i in range(inicio, fin + 1):
        with sr.AudioFile(f"{nombre_documento}{i}.wav") as source:
            audio_text = r.listen(source)

            try:
                text = r.recognize_google(audio_text, language="es-CL")
                print("Convirtiendo a texto")
                document = Document()
                document.add_paragraph(text)
                document.save(f"Transcripcion{i}.docx")
                print("La transcripción ha sido realizada con éxito")
                print("")
            except:
                print("Intenta de nuevo")

#6
def division_audio_partes_iguales():
    #Separa el audio en partes iguales de 3 minutos
    myaudio = input("Ingresa el nombre del audio con su extension. Ej: Clase1.wav: ")
    my_audio_1 = AudioFileClip(myaudio)
    print(f"La duración del video en segundos es de : {my_audio_1.duration} segundos")
    print(f"Puedes dividir el video en {my_audio_1.duration / 180} partes de 3 minutos cada una")
    cantidad_final = int(input("Redondea el número anterior hacia abajo: "))
    t_1 = 0
    t_2 = 180

    for i in range(1, cantidad_final + 1):
        clip_1 = my_audio_1.subclip(t_1, t_2)
        clip_1.write_audiofile(f"Parte {i}.wav")
        t_1 += 180
        t_2 += 180

    print("Se logró dividir tu audio de forma exitosa")

#7
fullText = []
def getText(filename):
    doc = docx.Document(filename)
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


#8
def loop_documentos():
    nombre_documento = input("Ingresa el nombre del documento sin la extensión: ")
    inicio = int(input("Ingresa el número del primer documento: "))
    fin = int(input("Ingresa el número del último documento: "))

    for i in range(inicio, fin+1):
        getText(f"{nombre_documento}{i}.docx")

    print("El loop fue completado")
#9
def juntar_textos_en_uno():
    documento_1 = Document()

    for i in fullText:
        documento_1.add_paragraph(i)

    documento_1.save("Textos_en_uno.docx")
    print("Tus documentos han sido juntados con éxito")
#10
def peticion_1():
    print("Que te gustaria hacer?: ")
    print(" 1) Descargar audio de un video de youtube")
    print(" 2) Descargar un video con su respectivo audio de youtube")
    print(" 3) Extraer el audio de un video que tú tengas")
    print(" 4) Convertir un video a un formato de audio")
    print(" 5) Ninguna de las anteriores")
    mensaje = input("Elige el número: ")

    if int(mensaje) == 1:
        descargar_audio()
    elif int(mensaje) == 2:
        descargar_video()
    elif int(mensaje) == 3:
        extraccion_audio()
    elif int(mensaje) == 4:
        conversor_video_a_audio()
    elif int(mensaje) == 5:
        pass

def peticion_2():
    print("¿Qué quieres hacer?")
    print("1 ) Dividir el audio en partes iguales(de un archivo de audio) y escribirlos como archivos separados")
    print("2 ) Ninguna de las anteriores")
    mensaje = input("Elige la opción: ")

    if int(mensaje) == 1:
        division_audio_partes_iguales()
    elif int(mensaje) == 2:
        pass

def peticion_3():
    print("¿Qué quieres hacer?")
    print("1) Convertir tus archivos de audio en texto")
    print("2) Ninguna de las anteriores")
    mensaje = int(input("Elige la opción: "))

    if mensaje == 1:
        audio_a_texto()
    elif mensaje == 2:
        pass


def peticion_4():
    print("¿Qué quieres hacer?")
    print("1) Juntar muchos documentos word en uno")
    print("2) Ninguna de las anteriores")
    mensaje = int(input("Elige la opción: "))

    if mensaje == 1:
        loop_documentos()
        print("Todavía no se juntan")
        juntar_textos_en_uno()
    elif mensaje == 2:
        pass